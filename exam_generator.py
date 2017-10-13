#!/usr/bin/env python
#
# Author:   Marcin Koculak <koculak.marcin@gmail.com>
# Created:  01.09.2017
#
# Copyright (C) Marcin Koculak
# For license information, see LICENSE.txt

import pandas as pd
import numpy as np
from string import ascii_lowercase


def shuffle_ans(series):
    '''
    Function shuffles answers for a question.
    
    First, it checks whether any answer should be left in place (indicated by %).
    Then it picks a new order for the remaining.
    '''
    mask = np.flatnonzero(~series.str.contains('%'))
    
    if len(mask)>0:
        series.iloc[mask] = np.random.permutation(series.iloc[mask])
    
    return series 


def random_exam(database, random = False, in_place = 0):
    '''
    Function generates a random exam from a given database (pandas Dataframe).
    Each row is one item. The first cell contains the question,
    while the answers are placed in subsequent cells.
    
    The correct answer must be indicated through adding a # (hash) sign
    at the end of the appropriate string.
    
    Additionally, one can add a % (percent) sign at the end of those answers
    that should preserve their place (e.g. being first or second answer).
    
    Parameters
    ----------
    random - boolean indicator whether answer should be shuffled
    
    in_place - integer parameter of how many answers from the end
               should stay in their place (e.g. for 'all of the above' type)
    
    '''
    
    rows, cols = database.shape
    
    # randomise the order of questions (rows)
    exam = database.sample(frac=1).reset_index(drop=True)
    
    if random:
        if in_place < 0:
            raise ValueError('in_place parameter cannot be smaller than 0')
        elif in_place <= cols - 1:
            # randomise the order of answers in questions
            exam.iloc[:,1:cols-in_place] = exam.iloc[:,1:cols-in_place].apply(shuffle_ans, axis=1)
        else:
            raise ValueError('in_place parameter cannot be grater than number of answers')
    
    # create a matrix with correct responses
    responses = []
    for n in range(rows):
        responses.append(np.flatnonzero(exam.iloc[n].str.contains('#'))[0])
    responses = pd.DataFrame(responses).applymap(lambda x: ascii_lowercase[x-1])
    
    # drop the unnecessary signs from the end of strings (hash, percent, space, dot)
    exam = exam.applymap(lambda x: x.rstrip('#% .'))
    return exam, responses


def format_exam(exam, answers, output_name = 'exam'):
    '''
    Fuction formats an exam and saves its in plain text.
    
    Questions are numbered and answers are listed with a), b) etc. markers.
    Questions are separated by an empty line.
    
    Parameters
    ----------
    exam - a Dataframe where each row is a question followed by its answers.
    
    answers - a Dataframe with indicators of correct answer to questions
                in a exam Dataframe.
    
    output_name - the name for the output files (exam and answers)
    '''
        
    rows, cols = exam.shape
    
    exam = ''
    for n in range(rows):
        question = str(n+1) + '. ' + question[0][n]+'\n'
        for m in range(1,cols):
            ans = '\t' + ascii_lowercase[m-1] + ') ' + question[m][n] + '.\n'
            question += ans
        exam += question
        exam += '\n'
    
    textfile = open(output_name + '.txt', 'w')
    textfile.write(exam)
    textfile.close()
    
    answers.to_excel(output_name+'.xlsx', header=False, index=False)


def format_exam_docx(exam, answers, output_name = 'exam', version = 1, template_file = 'template.docx', font='Lato'):
    '''
    Fuction formats an exam and saves its in MS Word format.
    
    Questions are numbered and answers are listed with a), b) etc. markers.
    Questions are separated by an empty line.
    
    Parameters
    ----------
    exam - a Dataframe where each row is a question followed by its answers.
    
    answers - a Dataframe with indicators of correct answer to questions
                in a exam Dataframe.
    
    output_name - the name for the output files (exam and answers).
    
    version - integer used to number the exam version.
    
    template_file - string indicating the location of the predefined docx template.
                    If empty, function uses a default unformated template.
    '''
    
    from docx import Document
    from docx.shared import Mm
    
    rows, cols = exam.shape
    
    if len(template_file):
        file = Document(template_file)
        zero = file.paragraphs[0]
    else:
        file = Document()
            
    for n in range(rows):
        question = file.add_paragraph('', style='List Number')
        question.add_run(exam.iloc[n,0]).font.name = font
        q_format = question.paragraph_format
        q_format.left_indent = Mm(7.5)
        q_format.first_line_indent = Mm(-7.5)
        for m in range(1,cols):
            answer = file.add_paragraph('', style='List')
            answer.add_run(ascii_lowercase[m-1] + ') ' + exam.iloc[n,m] + '.').font.name = font
            ans_format = answer.paragraph_format
            ans_format.left_indent = Mm(12)
            ans_format.first_line_indent = Mm(-4)
        
        empty_line = file.add_paragraph('')
    
    file.save(output_name + str(version) + '.docx')
    answers.to_excel(output_name + str(version) + '.xlsx', header=False, index=False)


base = pd.read_excel('exam.xlsx', header=None)
new, ans = random_exam(base, random=True, in_place=1)
format_exam_docx(new, ans, 'new_exam', version=1)
